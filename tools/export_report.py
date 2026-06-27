import base64
import io
import re
from typing import Optional
from urllib.parse import urlparse


def _links_to_domain(text: str) -> str:
    """Replace [label](url) with label (domain) for plain-text export."""
    def replace(m: re.Match) -> str:
        label, url = m.group(1), m.group(2)
        try:
            domain = urlparse(url).netloc.removeprefix("www.")
        except Exception:
            domain = ""
        return f"{label} ({domain})" if domain else label
    return re.sub(r"\[([^\]]+)\]\(([^\)]+)\)", replace, text)


def _generate_pdf(markdown_text: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_margins(left=15, top=15, right=15)
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    def safe(text: str) -> str:
        return text.encode("ascii", "ignore").decode("ascii")

    def write(size: int, bold: bool, text: str, line_height: float = 7.0) -> None:
        pdf.set_font("Helvetica", style="B" if bold else "", size=size)
        pdf.multi_cell(0, line_height, safe(text), new_x="LMARGIN", new_y="NEXT")

    for line in markdown_text.split("\n"):
        if line.startswith("# "):
            pdf.ln(2)
            write(18, True, line[2:], 10)
            pdf.ln(1)
        elif line.startswith("## "):
            pdf.ln(2)
            write(14, True, line[3:], 9)
            pdf.ln(1)
        elif line.startswith("### "):
            pdf.ln(1)
            write(12, True, line[4:], 8)
        elif line.startswith(("- ", "* ")):
            text = _links_to_domain(line[2:])
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            write(11, False, "- " + text)
        elif line.strip():
            text = _links_to_domain(line)
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            write(11, False, text)
        else:
            pdf.ln(3)

    return bytes(pdf.output())


def _generate_docx(markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in markdown_text.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith(("- ", "* ")):
            text = _links_to_domain(line[2:])
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
        elif line.strip():
            text = _links_to_domain(line)
            text = re.sub(r"\*+([^*]+)\*+", r"\1", text)
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_report(
    content: str,
    format: str,
    title: Optional[str] = None,
) -> dict:
    try:
        fmt = format.lower().strip()
        safe_title = (title or "report").replace(" ", "_")
        filename = f"{safe_title}.{fmt}"

        if fmt == "md":
            raw = content.encode("utf-8")
        elif fmt == "pdf":
            body = f"# {title}\n\n{content}" if title else content
            raw = _generate_pdf(body)
        elif fmt == "docx":
            body = f"# {title}\n\n{content}" if title else content
            raw = _generate_docx(body)
        else:
            return {
                "format": format,
                "filename": "",
                "content_base64": "",
                "success": False,
                "error": f"Unknown format '{format}'. Use pdf, docx, or md.",
            }

        return {
            "format": fmt,
            "filename": filename,
            "content_base64": base64.b64encode(raw).decode(),
            "success": True,
        }
    except Exception as exc:
        return {
            "format": format,
            "filename": "",
            "content_base64": "",
            "success": False,
            "error": str(exc),
        }
